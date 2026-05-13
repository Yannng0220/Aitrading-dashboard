from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("寶可夢對戰定位分類網站專案介紹")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("專案日期：2026-05-13\n").italic = True
    paragraph.add_run("GitHub Repo：https://github.com/Yannng0220/Aitrading-dashboard")


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbers(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_tech_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "模組"
    table.rows[0].cells[1].text = "使用技術"

    rows = [
        ("資料來源", "PokéAPI"),
        ("資料處理與訓練", "Python、Pandas、PyTorch、Torchvision"),
        ("模型架構", "ResNet18"),
        ("模型匯出", "ONNX"),
        ("瀏覽器推論", "ONNX Runtime Web"),
        ("網站前端", "HTML、CSS、JavaScript"),
        ("版本控管與部署", "GitHub、GitHub Pages、GitHub Actions"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right


def add_dataset_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "數量"
    table.rows[0].cells[2].text = "說明"

    rows = [
        ("總圖片數", "819", "所有寶可夢圖片總數"),
        ("attack", "239", "偏輸出型角色"),
        ("defense", "332", "偏坦克或高耐久角色"),
        ("support", "248", "偏功能型或輔助角色"),
        ("train", "655", "訓練集，約 80%"),
        ("test", "164", "測試集，約 20%"),
    ]
    for left, middle, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = middle
        row[2].text = right


def add_metrics_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["類別", "Precision", "Recall", "F1-score", "Support"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    rows = [
        ("attack", "0.5111", "0.4792", "0.4946", "48"),
        ("defense", "0.5600", "0.4242", "0.4828", "66"),
        ("support", "0.4638", "0.6400", "0.5378", "50"),
    ]
    for values in rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value


def add_file_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "檔案名稱"
    table.rows[0].cells[1].text = "用途說明"

    rows = [
        ("prepare_pokeapi_labels.py", "根據 PokéAPI 自動產生角色標籤與 train/test 切分"),
        ("dataset.py", "定義 PyTorch Dataset，負責載入圖片與標籤"),
        ("train.py", "模型訓練主程式，使用 ResNet18 完成三分類訓練"),
        ("predict.py", "單張圖片的本機預測程式"),
        ("export_web_model.py", "將 PyTorch 模型匯出成 ONNX 網頁模型"),
        ("index.html", "網站主頁"),
        ("web-app.js", "前端推論邏輯，負責讀圖、前處理、ONNX 推論與顯示結果"),
        ("styles.css", "網站版面與視覺樣式"),
        ("web/pokemon-role-classifier.onnx", "提供瀏覽器執行的模型檔"),
        ("web/model-config.json", "模型輸入大小、類別名稱、正規化參數設定"),
        ("labels.csv", "自動標註後的資料表"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right


def build_document() -> Document:
    document = Document()
    set_font(document)
    add_title(document)

    document.add_heading("一、專案簡介", level=1)
    document.add_paragraph(
        "本專案是一個結合 AI 影像分類與網頁部署的實作系統。使用者只要上傳一張寶可夢圖片，系統就會分析該角色在對戰中的定位，並將結果分類為攻擊（attack）、防禦（defense）或輔助（support）。"
    )
    document.add_paragraph(
        "本系統不只完成模型訓練，也完成從資料整理、自動標註、模型訓練、模型匯出到前端網站部署的完整流程，屬於一個可以直接展示的 AI 應用專案。"
    )

    document.add_heading("二、專案目標", level=1)
    add_bullets(
        document,
        [
            "建立一個能根據寶可夢圖片判斷其對戰定位的 AI 模型。",
            "利用 PokéAPI 自動建立資料標籤，降低人工標註成本。",
            "將模型轉換成瀏覽器可執行格式，讓網站不需要 Python 後端即可完成推論。",
            "把 AI 模型整合到實際網頁中，提升專案展示完整度。",
        ],
    )

    document.add_heading("三、系統功能", level=1)
    add_bullets(
        document,
        [
            "使用者可上傳 JPG、PNG、WebP 等圖片。",
            "網站會先顯示圖片預覽。",
            "系統會在瀏覽器中載入 ONNX 模型進行推論。",
            "輸出三種類別的機率值：attack、defense、support。",
            "顯示最可能的最終定位結果。",
        ],
    )

    document.add_heading("四、技術架構", level=1)
    add_tech_table(document)

    document.add_heading("五、資料集資訊", level=1)
    document.add_paragraph("本專案使用的寶可夢圖片資料集共有 819 張。資料經過自動標註後，分為三種類別。")
    add_dataset_table(document)

    document.add_heading("六、自動標註方法", level=1)
    document.add_paragraph(
        "由於原始圖片資料沒有直接標示對戰定位，因此本專案利用 PokéAPI 的資料進行自動標註。系統會根據圖片檔名解析出寶可夢的圖鑑編號與型態，再向 PokéAPI 取得對應資料。"
    )
    document.add_paragraph("自動標註時會參考的資訊包括：")
    add_bullets(document, ["種族值（stats）", "特性（abilities）", "屬性（types）", "招式列表（moves）"])
    document.add_paragraph(
        "系統接著透過規則加權計算 attack、defense、support 三種類型的分數，最後取最高分作為該圖片的標籤。這種方式能大幅降低人工逐張標註的時間。"
    )

    document.add_heading("七、模型訓練流程", level=1)
    add_numbers(
        document,
        [
            "讀取 labels.csv 與對應圖片檔案。",
            "將圖片縮放至 224×224，並做 Normalize 標準化。",
            "使用 ResNet18 作為影像分類模型。",
            "輸出三分類結果：attack、defense、support。",
            "訓練時使用 class-weighted loss 與 balanced sampler，降低類別不平衡影響。",
            "保存驗證表現最佳的模型權重 best_model.pt。",
        ],
    )

    document.add_heading("八、模型評估結果", level=1)
    document.add_paragraph("目前測試集的分類結果如下：")
    add_metrics_table(document)
    document.add_paragraph("整體 accuracy：0.5061（約 50.61%）")
    document.add_paragraph(
        "這代表模型已具備三分類能力，但仍有優化空間，例如增加資料量、調整標籤規則或導入更強的模型。"
    )

    document.add_heading("九、網站部署方式", level=1)
    document.add_paragraph(
        "為了讓模型可以直接在網站上執行，本專案將 PyTorch 訓練好的模型轉換成 ONNX 格式，並在前端使用 ONNX Runtime Web 進行推論。使用者上傳圖片後，網站會在本地瀏覽器完成模型運算，不需要 Python 後端。"
    )
    document.add_paragraph("網站推論流程如下：")
    add_numbers(
        document,
        [
            "使用者上傳圖片。",
            "前端將圖片縮放為 224×224。",
            "依照模型設定進行 Normalize。",
            "使用 ONNX Runtime Web 載入 ONNX 模型。",
            "輸出 logits，轉成 softmax 機率。",
            "顯示 attack、defense、support 機率與最終分類結果。",
        ],
    )

    document.add_heading("十、流程圖（文字版）", level=1)
    document.add_heading("1. 整體系統流程圖", level=2)
    document.add_paragraph(
        "寶可夢圖片資料集 → PokéAPI 自動標註 → labels.csv → ResNet18 訓練 → best_model.pt → 匯出 ONNX → 網頁載入模型 → 使用者上傳圖片 → 瀏覽器端推論 → 輸出 attack / defense / support"
    )
    document.add_heading("2. 自動標註流程圖", level=2)
    document.add_paragraph(
        "圖片檔名解析 → 取得圖鑑編號與型態 → 查詢 PokéAPI → 讀取 stats / abilities / types / moves → 規則加權評分 → 產生 attack / defense / support 分數 → 選取最高分 → 寫入 labels.csv"
    )
    document.add_heading("3. 模型訓練流程圖", level=2)
    document.add_paragraph("labels.csv + 圖片 → 影像前處理 → ResNet18 → 三分類輸出 → loss 計算 → 權重更新 → 保存最佳模型")
    document.add_heading("4. 網頁推論流程圖", level=2)
    document.add_paragraph("上傳圖片 → 顯示預覽 → 圖片縮放與標準化 → ONNX 模型推論 → softmax 機率 → 顯示最終定位與分數條")

    document.add_heading("十一、重要檔案說明", level=1)
    add_file_table(document)

    document.add_heading("十二、專案亮點", level=1)
    add_bullets(
        document,
        [
            "不是只有訓練模型，而是完成從資料、標籤、模型到網站部署的完整流程。",
            "使用 PokéAPI 自動標註，減少大量人工整理工作。",
            "模型可直接在瀏覽器端執行，不依賴 Python 後端。",
            "適合作為 AI 專案作品集、專題展示或剪報報告內容。",
        ],
    )

    document.add_heading("十三、目前限制", level=1)
    add_bullets(
        document,
        [
            "目前標籤是透過規則推論產生，不是官方直接提供的戰術定位。",
            "資料量仍偏小，圖片風格差異可能影響模型判斷。",
            "系統只看圖片，不看實際對戰隊伍搭配、個體值或努力值。",
            "ONNX 模型檔較大，網站首次載入時需要較多時間。",
        ],
    )

    document.add_heading("十四、未來改進方向", level=1)
    add_bullets(
        document,
        [
            "增加更多寶可夢圖片與不同畫風資料。",
            "優化標籤規則，提高 attack / defense / support 的區分能力。",
            "導入更強的影像模型架構。",
            "加入是否為寶可夢的前置辨識功能。",
            "加入更多對戰定位類別，例如速度型、坦克型、法術輸出型。",
        ],
    )

    document.add_heading("十五、簡報口頭介紹範例", level=1)
    document.add_paragraph(
        "本專案是一個寶可夢圖片定位分類系統。使用者只需要上傳一張寶可夢圖片，系統就會判斷它在對戰中偏向攻擊、防禦還是輔助。專案首先利用 PokéAPI 取得寶可夢的種族值、特性、屬性與招式資料，透過規則自動建立訓練標籤，再使用 ResNet18 模型進行影像分類訓練。訓練完成後，模型會被轉成 ONNX 格式，並部署到網頁前端，讓整個推論流程可以直接在瀏覽器中執行，不需要 Python 後端。這個專案的重點在於完成了從資料處理、模型訓練到網站部署的完整 AI 應用流程。"
    )

    document.add_heading("十六、附錄資訊", level=1)
    add_bullets(
        document,
        [
            "最新本地 commit：434a2d9 Fix GitHub Pages workflow setup",
            "前一版網站主 commit：0ef8c6a Deploy Pokemon role classifier web app",
            "GitHub Repo：https://github.com/Yannng0220/Aitrading-dashboard",
            "模型輸入尺寸：224 × 224",
            "模型輸出類別：attack、defense、support",
        ],
    )

    return document


def main() -> None:
    output_path = Path("Pokemon_Role_Classifier_Project_Report.docx")
    document = build_document()
    document.save(output_path)
    print(output_path.resolve())


if __name__ == "__main__":
    main()
